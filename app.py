from flask import Flask, render_template, request, send_file
from docx import Document
import pandas as pd
import re
import io
import os
from itertools import zip_longest
from openpyxl.utils import get_column_letter

app = Flask(__name__)

# Degrees (exact m/f forms)
DEGREES = {
    "ierindnieks","ierindniece",
    "kaprālis","kaprāliene",
    "seržants","seržante",
    "virsseržants","virsseržante",
    "virsniekvietnieks","virsniekvietniece",
    "leitnants","leitnante",
    "virsleitnants","virsleitnante",
    "kapteinis","kapteine",
    "majors","majore",
    "pulkvežleitnants","pulkvežleitnante",
    "pulkvedis","pulkvede",
    "ģenerālis","ģenerāle"
}

# Simple two-word name matcher
NAME_RE = re.compile(r'\b([A-ZĀČĒĢĪĶĻŅŖŠŪŽ][\w–]+)\s+([A-ZĀČĒĢĪĶĻŅŖŠŪŽ][\w–]+)\b')

# Helper patterns
NUM_ONLY = re.compile(r'^\d+\.\d+\.$')
DATE_RE = re.compile(r'202\d\. gada \d{1,2}\. [^\n,]+')
TIME_RE = re.compile(r'no plkst\.?\s*(\d{1,2}[:.]\d{2})\s*(?:līdz|–)\s*plkst\.?\s*(\d{1,2}[:.]\d{2})')

def extract_data(doc):
    # 1) Extract date & time
    full_text = "\n".join(p.text for p in doc.paragraphs)
    date_m = DATE_RE.search(full_text)
    time_m = TIME_RE.search(full_text)
    date_str = date_m.group().strip() if date_m else "N/A"
    time_str = f"{time_m.group(1)}–{time_m.group(2)}" if time_m else "N/A"
    full_dt = f"{date_str} {time_str}"

    # 2) Gather lines between "deleģētas" and the first "vadīs" after participants
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    start = next((i for i,t in enumerate(paras) if "deleģētas" in t), None)
    # We'll find the first occurrence of "vadīs" after start to mark the end of participants
    end = None
    if start is not None:
        for j in range(start+1, len(paras)):
            if "vadīs" in paras[j]:
                end = j
                break

    segment_lines = paras[start+1:end] if start is not None and end is not None else paras

    # Include every table cell text as separate lines
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    segment_lines.append(txt)

    # 3) Normalize into participant entries
    entries = []
    i = 0
    while i < len(segment_lines):
        ln = segment_lines[i]
        # If it's just a number like "1.1."
        if NUM_ONLY.match(ln):
            # next non-empty line is the real entry
            k = i+1
            while k < len(segment_lines) and not segment_lines[k].strip():
                k += 1
            if k < len(segment_lines):
                entries.append(segment_lines[k].strip())
            i = k + 1
            continue
        # If the line starts with one of the degree keywords
        first = ln.split()[0].lower()
        if first in DEGREES:
            entries.append(ln)
        i += 1

    # 4) Parse participants
    participants = []
    for seg in entries:
        parts = seg.split()
        deg = parts[0] if parts and parts[0].lower() in DEGREES else ""
        remainder = seg[len(deg):].strip() if deg else seg

        # Extract name by NAME_RE
        name = ""
        nm = NAME_RE.search(remainder)
        if nm:
            name = f"{nm.group(1)} {nm.group(2)}"

        # Extract job as everything after the comma following the name
        job = ""
        if ',' in remainder and name:
            # safe split on f"{name},"
            _,_,after = remainder.partition(f"{name},")
            job = after.strip()

        participants.append({
            "degree": deg,
            "participant": name,
            "pjob": job
        })

    # Identify the paragraph index of the last participant (for lecturer scan)
    last_part_para = None
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if any(txt.startswith(f"{n}.") for n in range(1,10)):
            last_part_para = idx

    # 5) Extract lecturers: find **all** paragraphs after last_part_para containing "vadīs"
    lecturers = []
    start_idx = last_part_para+1 if last_part_para is not None else 0
    for p in doc.paragraphs[start_idx:]:
        txt = p.text.strip()
        if "vadīs" not in txt:
            continue
        # Tail after the first "vadīs"
        parts = re.split(r'vadīs[:\-]?', txt, 1, flags=re.IGNORECASE)
        tail = parts[1] if len(parts) > 1 else parts[0]
        # Split multiple lecturers by semicolon or the word " un "
        segments = re.split(r';|\s+un\s+', tail)
        for seg in segments:
            entry = seg.strip().rstrip('.;')
            # Extract name
            nm = ""
            jb = ""
            m = NAME_RE.search(entry)
            if m:
                nm = f"{m.group(1)} {m.group(2)}"
                # Extract job after name comma
                _,sep,suff = entry.partition(f"{nm},")
                jb = suff.strip() if sep else ""
            else:
                # fallback: split on first comma
                first,sep,rest = entry.partition(',')
                nm = first.strip()
                jb = rest.strip() if sep else ""
            if nm or jb:
                lecturers.append({
                    "lecturer": nm,
                    "ljob": jb
                })

    # 6) Build rows by zipping participants and lecturers
    rows = []
    for i, (p, l) in enumerate(zip_longest(participants, lecturers, fillvalue={})):
        rows.append({
            "Date": full_dt if i == 0 else "",
            "Degree": p.get("degree", ""),
            "Participant": p.get("participant", ""),
            "Participant Job": p.get("pjob", ""),
            "Lecturer": l.get("lecturer", ""),
            "Lecturer Job": l.get("ljob", ""),
        })

    df = pd.DataFrame(rows)

    # 7) Write Excel with auto-fit columns
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Data")
        ws = writer.sheets["Data"]
        for ci, col in enumerate(df.columns, 1):
            maxlen = df[col].astype(str).map(len).max()
            ws.column_dimensions[get_column_letter(ci)].width = max(maxlen, len(col)) + 2
    out.seek(0)
    return out

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/upload", methods=["POST"])
def upload():
    f = request.files.get("file")
    if not f or not f.filename.lower().endswith(".docx"):
        return "Please upload a .docx file", 400
    doc = Document(f)
    excel_io = extract_data(doc)
    return send_file(
        excel_io,
        download_name="seminar_data.xlsx",
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 10000)), debug=True)
